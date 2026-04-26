"""Generate the exhibition walkthrough Word document.

Run:
    python scripts/build_exhibition_walkthrough.py
Output:
    docs/ICBB_AMR_System_Exhibition_Walkthrough.docx
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "docs" / "ICBB_AMR_System_Exhibition_Walkthrough.docx"

FOREST = RGBColor(0x1D, 0x4D, 0x43)
GOLD   = RGBColor(0xB1, 0x8A, 0x3A)
INK    = RGBColor(0x1A, 0x17, 0x0E)


def _set_base_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = INK


def H(doc, text, level=1, color=FOREST):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return p


def P(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = INK
    return p


def BULLET(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = INK
    return p


def NUM(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = INK
    return p


def QUOTE(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = FOREST


def DEFINITION(doc, term, meaning):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{term} — ")
    r1.bold = True
    r1.font.color.rgb = FOREST
    r2 = p.add_run(meaning)
    r2.font.color.rgb = INK


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_base_font(doc)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("ICBB AMR Surveillance System")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = FOREST

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Exhibition Walkthrough · A Plain-Language Guide for Visitors, Decision-Makers and Investors")
    sr.italic = True
    sr.font.size = Pt(13)
    sr.font.color.rgb = GOLD

    P(doc, "")
    P(doc, "Prepared for the National AMR Exhibition · Ghana", italic=True, size=10)
    P(doc, "A unified, One-Health digital platform for tracking antimicrobial resistance across the environment, food, human, animal and aquaculture sectors.", italic=True, size=11)
    doc.add_page_break()

    # 1. The big picture
    H(doc, "1. The Big Picture — What This System Is", level=1)
    P(doc,
      "The ICBB AMR Surveillance System is Ghana’s national digital platform for tracking "
      "antimicrobial resistance (AMR). In simple terms, it is a website that any approved laboratory "
      "in Ghana can log into to upload its laboratory test results, and that public-health officials, "
      "veterinarians, food-safety officers and policy-makers can use to see — in real time — where "
      "drug-resistant infections are spreading, which medicines are still working, and which are no "
      "longer effective.")
    P(doc,
      "It brings together five worlds that have always been managed separately:")
    BULLET(doc, "Human health (hospitals, clinics)")
    BULLET(doc, "Animal health (veterinary laboratories, farms)")
    BULLET(doc, "Food (slaughterhouses, markets, processors)")
    BULLET(doc, "Environment (rivers, wastewater, soil)")
    BULLET(doc, "Aquaculture (fish farms, ponds)")
    P(doc,
      "This is the One-Health approach: bacteria do not respect the boundary between a hospital ward "
      "and a chicken farm, so the surveillance has to look at all five worlds at the same time.")

    # 2. What is AMR
    H(doc, "2. What Is Antimicrobial Resistance, in One Paragraph?", level=1)
    P(doc,
      "When a person, animal or fish is treated repeatedly with antibiotics, the bacteria that "
      "survive the treatment learn to defend themselves. Over time, those bacteria multiply, spread "
      "to other people, animals and into the environment, and the antibiotic stops working. A simple "
      "wound infection or a urinary tract infection can then become deadly because none of the "
      "available drugs can kill the bacteria. AMR is already estimated to kill more people each year "
      "than malaria and HIV combined, and Ghana — like every other country — has no choice but to "
      "monitor it carefully and act early.")

    # 3. Walkthrough
    H(doc, "3. Live Walkthrough — A Page-by-Page Tour", level=1)
    P(doc,
      "The system has 21 working pages grouped into seven simple zones. As you walk visitors through "
      "the demonstration, follow this order — it tells a complete story from raw data to national "
      "decision-making.")

    H(doc, "3.1 Secure Login", level=2)
    P(doc,
      "Every visitor sees a clean two-pane login screen. On the left is the national branding; on "
      "the right, an email and password field. Only the 19 approved sentinel laboratories and the "
      "national administrator can sign in. Passwords are stored using bcrypt encryption (the same "
      "standard banks use), and the session times out automatically when idle. The login also "
      "decides what the user is allowed to see: a hospital lab will only see its own data, while "
      "the national administrator sees everything.")
    QUOTE(doc, "Key message for visitors: ‘The system protects each laboratory’s data while still "
                "letting the country see the national picture.’")

    H(doc, "3.2 Zone A — Data In (Upload & Quality)", level=2)
    P(doc,
      "Laboratories enter their data into the system in one of two equally supported ways. The "
      "choice is the lab’s — nothing on the analysis side changes.",)
    DEFINITION(doc, "Option 1 — KoboToolbox form (recommended)",
               "The national programme has published a single standard AMR data-collection form on "
               "KoboToolbox. Each sentinel laboratory fills it in on a phone, tablet or laptop — "
               "online or offline — and submits. The dashboard then pulls every new submission into "
               "the national database simply by knowing the form’s ID. No file handling, no email "
               "attachments, no version-control problems.")
    DEFINITION(doc, "Option 2 — Excel upload",
               "Labs that already keep their records in Excel can upload the workbook directly — "
               "useful for legacy data and for partners who are not yet on KoboToolbox.")
    DEFINITION(doc, "Data Management",
               "Lets each lab review, rename or delete its own datasets. The administrator can "
               "additionally promote one dataset to be the ‘National Main Dataset’ — that is the "
               "one used in all the national charts.")
    P(doc,
      "Behind the scenes, whichever route is used, the system standardises pathogen names, "
      "antibiotic codes and specimen types using internationally accepted look-up tables, so that "
      "a laboratory in Tamale and one in Cape Coast can be compared on equal terms. WHONET "
      "appears later in the workflow only as an export format — the system can re-emit cleaned "
      "data in WHONET form so it can be shared with WHO GLASS or re-imported into the desktop "
      "WHONET software.")

    H(doc, "3.3 Zone B — Surveillance Dashboards (the heart of the demo)", level=2)
    DEFINITION(doc, "Resistance Overview",
               "The headline page. Big number cards show: total isolates, % multi-drug-resistant "
               "(MDR), top 5 priority pathogens, and the share of last-line antibiotics that have "
               "already failed. Visitors instantly grasp the scale.")
    DEFINITION(doc, "Resistance Heat Map",
               "A coloured grid where each row is a bacterium and each column is an antibiotic. "
               "Green = the drug still works, red = it has failed. One glance tells a clinician "
               "what to prescribe.")
    DEFINITION(doc, "Pathogen Profile",
               "Pick any organism (e.g. E. coli) and see its full resistance fingerprint, where it "
               "has been found in the country, in which sector, and how its resistance has changed "
               "over time.")
    DEFINITION(doc, "HAI Profile",
               "Healthcare-Associated Infections: focuses on the bacteria patients catch inside "
               "the hospital itself — a leading cause of preventable death.")
    DEFINITION(doc, "Trends",
               "Time-series line charts showing whether resistance to each antibiotic is rising, "
               "falling or stable, with monthly / quarterly / yearly aggregation and earliest / "
               "latest test markers.")
    DEFINITION(doc, "Antibiogram Matrix",
               "The clinician’s cheat-sheet: a table of % susceptibility for every "
               "pathogen-antibiotic pair, colour-coded:")
    BULLET(doc, "Green ≥ 90 %  — drug still highly effective")
    BULLET(doc, "Lime 70–89 % — usually effective")
    BULLET(doc, "Yellow 50–69 % — caution; consider alternatives")
    BULLET(doc, "Orange 30–49 % — likely to fail")
    BULLET(doc, "Red < 30 % — do not use empirically")
    P(doc, "An asterisk (*) marks pathogen-drug pairs with fewer than the minimum CLSI-recommended "
           "30 isolates — a built-in honesty signal so clinicians know when to interpret the cell with caution.")

    H(doc, "3.4 Zone C — Geography (where is the problem?)", level=2)
    DEFINITION(doc, "Map Hotspots",
               "An interactive map of Ghana. Each region or sentinel site is plotted with a coloured "
               "circle whose size = number of isolates and whose colour = resistance rate. Visitors "
               "can filter by sector, year and pathogen and immediately see, for example, that "
               "carbapenem-resistant E. coli is concentrated in two specific districts.")

    H(doc, "3.5 Zone D — Intelligence (what is going to happen?)", level=2)
    DEFINITION(doc, "Advanced Analytics",
               "Statistical breakdowns: resistance by age group, sex, ward, animal species, water "
               "source, etc. Helps researchers find risk factors.")
    DEFINITION(doc, "Risk Assessment",
               "A scoring engine that combines pathogen virulence, current resistance rate, "
               "spread potential and clinical impact into a single 0–100 risk score per "
               "pathogen-region pair, then ranks them from highest to lowest priority.")
    DEFINITION(doc, "Comparative Analysis",
               "Side-by-side comparison of any two regions, two laboratories, two years or two "
               "sectors. Used to answer questions like ‘Has resistance in poultry farms gotten worse "
               "than in hospitals over the last 3 years?’")
    DEFINITION(doc, "Alerts Dashboard",
               "Scans the latest data against user-configurable severity thresholds (Critical / "
               "High / Medium / Low) and displays a colour-coded summary plus a list of every "
               "flagged event — for example, three or more carbapenem-resistant isolates from the "
               "same hospital in one week. The alert is shown in the dashboard whenever a user "
               "opens this page; automatic email/SMS delivery is on the funded roadmap (see §11).")

    H(doc, "3.6 Zone E — Use & Consumption (the other half of AMR)", level=2)
    P(doc,
      "Resistance is only half the story. The other half is how much antibiotic is being used. "
      "The system tracks this in three dedicated dashboards:")
    DEFINITION(doc, "PPS Dashboard",
               "Point-Prevalence Surveys done in hospitals on a chosen day, recording every patient "
               "currently on an antibiotic — the gold-standard method for measuring hospital "
               "prescribing patterns.")
    DEFINITION(doc, "AMU Dashboard",
               "Antimicrobial Use, expressed in DDD (Defined Daily Doses) per 1,000 patient-days. "
               "DDD is the WHO’s international yard-stick that lets you compare a 200-bed district "
               "hospital with a 1,500-bed teaching hospital fairly.")
    DEFINITION(doc, "AMC Dashboard",
               "Antimicrobial Consumption — measured at the country level using import, "
               "manufacturing and sales data, expressed in DDD per 1,000 inhabitants per day. This "
               "is the figure Ghana reports to the WHO GLASS programme.")

    H(doc, "3.7 Zone F — Reporting & Export", level=2)
    DEFINITION(doc, "Antibiogram (PDF/Word)",
               "Generates a print-ready antibiogram for any laboratory or for the country, ready to "
               "post on a hospital ward or to ship to WHO.")
    DEFINITION(doc, "WHONET Export",
               "Re-exports the cleaned data in WHONET format so it can be shared with international "
               "partners or re-imported into the desktop WHONET software.")
    DEFINITION(doc, "Report Export",
               "One-click national technical surveillance report (PDF or DOCX) with an executive "
               "summary, every chart, every table and a list of recommended actions — what would "
               "normally take a team of analysts three weeks is produced in 90 seconds.")

    H(doc, "3.8 Zone G — Administration (national focal point only)", level=2)
    BULLET(doc, "Admin – Datasets: choose the National Main Dataset, archive obsolete uploads.")
    BULLET(doc, "Admin – Users: approve new laboratories, reset passwords, deactivate leavers.")

    # 4. Variables explained
    H(doc, "4. The Numbers Explained — Every Variable in Plain Language", level=1)
    P(doc,
      "Below is a glossary of every metric a visitor will see on screen. Use it as your cheat-sheet "
      "during the demo.")

    items = [
        ("Isolate",
         "One bacterium grown from one specimen from one patient/animal/sample. The basic unit of counting."),
        ("Sample / Specimen",
         "The clinical material the bacterium came from — urine, blood, stool, wound swab, milk, "
         "carcass swab, water, etc."),
        ("AST (Antibiotic Susceptibility Test)",
         "The lab test that exposes a bacterium to an antibiotic and reports whether it is "
         "Susceptible (S), Intermediate (I) or Resistant (R)."),
        ("% Susceptible",
         "Of all isolates of a given bacterium tested against a given antibiotic, the share that "
         "were S. Calculated as: (S ÷ (S + I + R)) × 100. Higher = the drug still works."),
        ("% Resistant",
         "The opposite of % Susceptible: (R ÷ (S + I + R)) × 100. This is the headline AMR figure."),
        ("MDR (Multi-Drug-Resistant)",
         "An isolate that is non-susceptible to at least one agent in three or more antibiotic "
         "classes — the international Magiorakos definition."),
        ("XDR (Extensively Drug-Resistant)",
         "Non-susceptible to all but two or fewer classes — only a handful of treatment options remain."),
        ("PDR (Pan-Drug-Resistant)",
         "Resistant to every antibiotic tested — effectively untreatable with current drugs."),
        ("Priority Pathogens",
         "The WHO list of bacteria of greatest public-health concern: Critical (e.g. carbapenem-resistant "
         "Acinetobacter, E. coli, Klebsiella), High (e.g. methicillin-resistant Staphylococcus aureus / MRSA, "
         "vancomycin-resistant Enterococci), Medium."),
        ("ESBL",
         "Extended-Spectrum Beta-Lactamase: an enzyme some bacteria make that destroys most penicillins "
         "and cephalosporins. The system flags ESBL-producers automatically."),
        ("CRE / CRAB",
         "Carbapenem-Resistant Enterobacterales / Carbapenem-Resistant Acinetobacter baumannii — "
         "WHO Critical-priority threats. When the dashboard shows even a single CRE in a sentinel "
         "site, it is treated as a red-alert event."),
        ("DDD (Defined Daily Dose)",
         "The WHO-assigned average maintenance dose per day for a drug used for its main indication "
         "in adults. Lets us compare a 250 mg tablet of ciprofloxacin with a 1 g vial of ceftriaxone "
         "fairly."),
        ("DDD per 1,000 inhabitants per day",
         "The standard national antibiotic-consumption metric used by GLASS and ESAC-Net. A figure "
         "of ‘15 DDD/1000/day’ means that on any given day, 15 out of every 1,000 Ghanaians is "
         "taking a defined daily dose of antibiotics."),
        ("DDD per 100 patient-days",
         "The hospital-level use metric. Independent of bed size, so a small district hospital and a "
         "teaching hospital can be compared."),
        ("Point-Prevalence (PPS)",
         "A snapshot survey done on one day: the % of patients in the hospital who are on at least "
         "one antibiotic at that moment."),
        ("AWaRe Categories",
         "WHO’s classification of antibiotics into Access (first-line, low resistance risk), Watch "
         "(higher risk, use sparingly) and Reserve (last-line, only when nothing else works). The "
         "national target is ≥ 60 % of all use coming from the Access group."),
        ("Sector",
         "Human / Animal / Food / Environment / Aquaculture. Every isolate carries its sector tag, "
         "which is what makes One-Health analytics possible."),
        ("Sentinel Site",
         "An approved laboratory that has been trained, audited and connected to the national "
         "system. Currently 19 sites across all 16 regions of Ghana."),
        ("Confidence Interval",
         "A statistical range expressing how reliable a percentage is given the number of isolates "
         "behind it. The current system enforces this honestly through the CLSI minimum-isolate "
         "rule (cells based on fewer than the threshold are flagged with an asterisk); displaying "
         "explicit numerical CI bands on the on-screen Trends chart is on the roadmap (§11)."),
        ("Risk Score (0–100)",
         "Composite indicator combining (a) current resistance rate, (b) trend slope, (c) WHO "
         "priority class, (d) sample volume, (e) clinical impact. Used to rank where intervention "
         "money will save the most lives."),
    ]
    for term, meaning in items:
        DEFINITION(doc, term, meaning)

    # 5. How it works under the hood
    H(doc, "5. How It Works Under the Hood (in non-technical terms)", level=1)
    NUM(doc, "A laboratory either fills the standard KoboToolbox AMR form on a phone / tablet / "
             "computer (online or offline), or prepares an Excel workbook in the agreed template.")
    NUM(doc, "For KoboToolbox: the national dashboard pulls every new submission automatically by "
             "the form’s ID. For Excel: the lab signs in to the website and uploads the workbook. "
             "Nothing is installed; any computer or tablet works.")
    NUM(doc, "An automatic quality-control engine cleans the data: it harmonises pathogen names, "
             "drops duplicate isolates from the same patient, flags impossible values and assigns a "
             "data-quality score.")
    NUM(doc, "Clean data flows into a secure cloud database (PostgreSQL on Supabase, hosted in EU "
             "data-centres with daily back-ups and TLS encryption end-to-end).")
    NUM(doc, "Every dashboard re-reads the data on demand, applies the filters the user has chosen, "
             "computes the resistance percentages, MDR rates, DDDs and risk scores, and renders "
             "interactive charts.")
    NUM(doc, "The Alerts Dashboard generates threshold-breach events whenever a user opens it. "
             "Automatic 24/7 monitoring with email/SMS notification is implemented in code but "
             "currently switched off pending funded re-deployment.")
    NUM(doc, "National technical surveillance reports can be generated on demand at any time; "
             "automatic monthly/quarterly scheduling is on the roadmap.")
    P(doc,
      "The platform is built with open, free-to-use technology (Python, Streamlit, PostgreSQL, "
      "Plotly, Folium) — no licence fees, no vendor lock-in. Source code is version-controlled on "
      "GitHub and the running app is hosted on Streamlit Community Cloud, meaning the marginal "
      "cost of adding one more sentinel site is essentially zero.")

    # 6. Who benefits
    H(doc, "6. Who Benefits, and How", level=1)
    DEFINITION(doc, "Clinicians", "Get a live, evidence-based antibiogram so they stop prescribing "
                                  "drugs that have already failed in their region.")
    DEFINITION(doc, "Hospital pharmacists", "See real-time consumption vs. resistance — the basis "
                                            "of every Antimicrobial Stewardship Programme.")
    DEFINITION(doc, "Veterinarians & farmers", "Discover which farms are amplifying resistance and "
                                               "can adjust husbandry, vaccination and antibiotic "
                                               "use accordingly.")
    DEFINITION(doc, "Food-safety inspectors", "Trace resistant strains from the farm to the market "
                                              "to the kitchen.")
    DEFINITION(doc, "Environmental scientists", "Map resistance hotspots in rivers and wastewater "
                                                "and link them to upstream sources.")
    DEFINITION(doc, "Policy-makers & MoH/MoFA", "Get a single, credible national figure for "
                                                "international reporting (WHO GLASS, OIE, FAO, "
                                                "WOAH) and for budget allocation.")
    DEFINITION(doc, "The general public", "Lower risk of dying from a treatable infection; cheaper "
                                          "drugs because we keep the older ones working; safer food, "
                                          "water and meat.")

    # 7. Why invest
    H(doc, "7. The Investment Case — Why Ghana Should Fund This System", level=1)
    P(doc,
      "Antimicrobial resistance is projected by the World Bank to push 28 million more Africans "
      "into extreme poverty by 2050 and to cost the continent up to 5 % of its GDP every year if "
      "left unchecked. Ghana cannot afford either of those numbers.", italic=True)
    P(doc,
      "The ICBB AMR Surveillance System is the cheapest insurance policy the country can buy, for "
      "five reasons:")
    NUM(doc, "It already exists and already works. Investment is needed to scale it, not to build "
             "it from scratch — the highest-risk and most expensive phase is behind us.")
    NUM(doc, "It uses entirely open-source technology, so 100 % of every cedi invested goes into "
             "people, training, sentinel sites and data quality — none into software licences.")
    NUM(doc, "One national platform replaces 19 separate Excel-based reports and several donor-funded "
             "parallel systems. The savings in duplicated effort alone pay for the running cost.")
    NUM(doc, "It produces the exact indicators required by WHO GLASS, the African Union AMR "
             "framework and the Tripartite (WHO/FAO/WOAH) One-Health agenda — so each report "
             "Ghana submits is automatically generated, audit-ready and internationally comparable.")
    NUM(doc, "Every cedi spent on AMR surveillance has been estimated by the World Bank and the "
             "O’Neill Review on AMR (2016) to return many times its value in averted treatment costs, "
             "shorter hospital stays and avoided deaths. A Ghana-specific cost-benefit study is one "
             "of the deliverables of the next funding phase.")
    P(doc, "What investment unlocks:", bold=True)
    BULLET(doc, "Expansion from 19 to 50+ sentinel sites covering every district capital.")
    BULLET(doc, "On-site WHONET training and data-quality auditing for every site, every year.")
    BULLET(doc, "Integration with veterinary, food and environmental laboratories that are not yet connected.")
    BULLET(doc, "An SMS / mobile interface so rural health workers can submit data from a basic phone.")
    BULLET(doc, "A national AMR data warehouse that researchers and African partners can query "
                "(with appropriate governance).")
    BULLET(doc, "Long-term cloud hosting on a sovereign Ghanaian data centre.")
    P(doc,
      "Funding the ICBB AMR Surveillance System is not just funding an IT platform — it is funding "
      "the country’s ability to keep antibiotics working for our children. The drugs we are losing "
      "today took the world 70 years to discover. Without surveillance we will lose them silently. "
      "With surveillance we can fight back.", bold=True)

    # 8. A voice to the public
    H(doc, "8. A Voice to the Public", level=1)
    QUOTE(doc, "Antibiotics are a shared inheritance. Every time one of us takes them when we don’t "
                "need to, or stops them too soon, or buys them without prescription, we use up a "
                "little of that inheritance for everyone. The ICBB AMR Surveillance System is "
                "Ghana’s way of measuring how much of the inheritance is left — and of giving it "
                "back to our children.")
    P(doc, "Three things every Ghanaian can do today:")
    NUM(doc, "Only take antibiotics when prescribed by a qualified health worker — and finish the full course.")
    NUM(doc, "Never share leftover antibiotics, and never buy them ‘over the counter’ from unlicensed sellers.")
    NUM(doc, "Vaccinate yourself, your children and your animals — every infection prevented is an antibiotic course saved.")

    # 9. Demo script
    H(doc, "9. A 5-Minute Demo Script for the Booth", level=1)
    NUM(doc, "‘This is Ghana’s national AMR surveillance platform. Let me sign in as a hospital lab.’ — show login.")
    NUM(doc, "‘Here we’ve uploaded last quarter’s WHONET file. The system flagged 3 % of rows as questionable — those are highlighted.’ — Upload page.")
    NUM(doc, "‘Now let’s switch to the national view.’ — Resistance Overview: read out total isolates, % MDR, top 3 pathogens.")
    NUM(doc, "‘This heat-map is what every clinician really wants.’ — Heat Map: point at one red cell. ‘That drug is finished for that bug in this country.’")
    NUM(doc, "‘And this is where it’s happening.’ — Map Hotspots: zoom into one region.")
    NUM(doc, "‘The system is also watching for outbreaks.’ — Alerts Dashboard: show one open alert.")
    NUM(doc, "‘When the Minister asks for the quarterly report, this is what used to take three weeks.’ — Report Export, click, show the generated PDF.")
    NUM(doc, "Close with the investment ask: ‘We have built it; we now need to scale it. Will you help us reach all 50 districts?’")

    # 10. FAQ
    H(doc, "10. Anticipated Questions & Short Answers", level=1)
    DEFINITION(doc, "Is patient data safe?",
               "Yes. We never store patient names, addresses or NHIS numbers. Only de-identified "
               "isolate-level data, encrypted in transit and at rest. Each lab can only see its own "
               "patients.")
    DEFINITION(doc, "What if a lab has no internet?",
               "Labs can prepare WHONET files offline and upload in batches when connectivity is "
               "available. An SMS submission gateway is on the roadmap.")
    DEFINITION(doc, "How is this different from the WHO GLASS portal?",
               "GLASS is the global aggregator. This system is what feeds GLASS — and adds the "
               "veterinary, food and environmental sectors that GLASS does not yet cover.")
    DEFINITION(doc, "Who owns the data?",
               "The Government of Ghana, through the ICBB programme. Sentinel laboratories own "
               "their submissions and decide on data-sharing agreements.")
    DEFINITION(doc, "How accurate are the percentages?",
               "Every figure comes with an isolate count and the CLSI minimum-isolate rule is "
               "enforced. Cells based on fewer than 30 isolates are flagged with an asterisk so no "
               "one over-interprets a small sample. Explicit numerical confidence-interval bands "
               "on screen are on the roadmap (§11).")

    # 11. Roadmap
    H(doc, "11. Roadmap — What the Next Funded Phase Will Add", level=1)
    P(doc,
      "The platform shown today is the working core. The features below are designed, partially "
      "implemented or in code but switched off, and will be activated as the next funding phase "
      "lands. They are listed here so investors can see exactly what their support will unlock.",)

    DEFINITION(doc, "Automated email & SMS alerts",
               "The Alerts Dashboard already detects threshold breaches in-app. The notification "
               "layer (SMTP / SendGrid / SMS gateway) was built and is preserved in the code base "
               "but is currently disconnected. Re-enabling it will mean the National Focal Point "
               "and the responsible lab receive an email or SMS within minutes of a Critical or "
               "High-priority event.")
    DEFINITION(doc, "24/7 background alert watcher",
               "Today, alerts are computed when a user opens the Alerts page. A scheduled "
               "background worker will scan the incoming data continuously so events are detected "
               "even when nobody is logged in.")
    DEFINITION(doc, "Scheduled automatic national reports",
               "Reports are currently generated on demand (one click → PDF/Word in ~90 seconds). "
               "The next phase will add a scheduler that produces and emails the monthly and "
               "quarterly national surveillance reports automatically to a configured distribution "
               "list.")
    DEFINITION(doc, "Confidence-interval bands on Trends charts",
               "The Trends page currently shows the resistance trajectory; explicit numerical 95 % "
               "confidence-interval bands (e.g. Wilson or Clopper–Pearson) will be added so users "
               "can see at a glance whether a change is statistically real.")
    DEFINITION(doc, "Forecasting on every dashboard",
               "A short-horizon forecast engine already exists in the report module; surfacing it "
               "as an on-screen 6-month projection on the Pathogen and Trends pages is the next step.")
    DEFINITION(doc, "Mobile / SMS data submission for low-bandwidth sites",
               "An SMS-to-database gateway so rural health workers can submit isolate counts from "
               "a basic phone when KoboToolbox or web upload is impractical.")
    DEFINITION(doc, "Direct LIS integration",
               "Live, automated pull from hospital Laboratory Information Systems (where they "
               "exist), removing the manual step of filling the KoboToolbox form.")
    DEFINITION(doc, "Sentinel-site expansion (19 → 50+)",
               "On-site WHONET / KoboToolbox training and data-quality auditing for every "
               "additional district hospital and reference laboratory.")
    DEFINITION(doc, "Veterinary, food and environmental sector deepening",
               "More routine sampling streams from poultry, livestock, aquaculture, slaughterhouses, "
               "rivers and wastewater to make the One-Health picture truly representative.")
    DEFINITION(doc, "Sovereign Ghanaian hosting",
               "Migration from the current EU-region cloud provider to a Ghana-based government data "
               "centre once one is procured, for full data sovereignty.")
    DEFINITION(doc, "Public dashboard",
               "A read-only, aggregated public view of national resistance trends and AWaRe "
               "consumption, accessible to journalists, schools and the general public without "
               "exposing any patient or isolate-level data.")
    DEFINITION(doc, "Ghana-specific cost-benefit evaluation",
               "An independent study quantifying the cedi-for-cedi return on investment in AMR "
               "surveillance under Ghanaian unit costs, to complement the global O’Neill estimates.")

    # Footer
    P(doc, "")
    P(doc, "— End of Walkthrough —", italic=True, size=10)
    P(doc, "Prepared by the ICBB AMR Surveillance Team. For more information, contact the National "
           "AMR Focal Point. © ICBB · National AMR Surveillance Programme.", italic=True, size=9)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
