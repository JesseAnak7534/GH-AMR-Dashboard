"""
Provision login credentials for every approved sentinel-site laboratory.

For each lab in `src.lab_management.APPROVED_LABS` this script:

  1. Assigns a short, easy-to-type username (email) of the form
        <code>@icbb-amr.gh
     where <code> is a curated short acronym for the lab.

  2. Generates a stable, memorable but reasonably strong password of the form
        <CODE>-Amr@<4-random-digits>
     e.g. ``KBTH-Amr@4827``  (>= 12 chars, mixed case, digit, symbol).

  3. Persists the plaintext credentials to ``db/lab_credentials.json`` so the
     same passwords are reused on subsequent runs (idempotent).  The file is
     gitignored.

  4. Inserts the user into the ``users`` table using bcrypt.  If the user
     already exists, the password is reset to the persisted value so the
     credentials shared with the lab keep working.

  5. Writes a Word document ``Lab_Login_Credentials.docx`` (one row per lab)
     for the administrator to share with sentinel sites.

These credentials are intentionally **fixed** -- the labs use them for live
data submission until the administrator chooses to rotate them.

Usage:
    .\\venv\\Scripts\\python.exe scripts\\setup_lab_logins.py
"""

from __future__ import annotations

import json
import os
import secrets
import sys
from pathlib import Path

import bcrypt

# Allow running as a top-level script.
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src import db  # noqa: E402
from src.lab_management import (  # noqa: E402
    APPROVED_LABS,
    LAB_EMAIL_DOMAIN as _LAB_EMAIL_DOMAIN,
    LAB_LOGIN_CODES,
)


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

# Single source of truth lives in src.lab_management so the dashboard's
# sign-in flow and this provisioner can never drift apart.
EMAIL_DOMAIN = _LAB_EMAIL_DOMAIN
LAB_CODES: dict[str, str] = dict(LAB_LOGIN_CODES)
CREDENTIALS_FILE = Path("db") / "lab_credentials.json"
# Public, hash-only manifest committed to the repo so the deployed Streamlit
# Cloud app can bootstrap the same lab accounts into its own database on
# startup.  Plaintext passwords are NEVER written here.
PUBLIC_HASH_FILE = Path("db") / "lab_accounts.json"
WORD_OUTPUT = Path("Lab_Login_Credentials.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_password(code: str) -> str:
    """Return a memorable but reasonably strong password.

    Pattern: ``<CODE>-Amr@<4 digits>`` -- always >= 11 chars and contains
    upper-case, lower-case, digit and a symbol.
    """
    digits = "".join(secrets.choice("0123456789") for _ in range(4))
    return f"{code.upper()}-Amr@{digits}"


def _load_existing() -> dict:
    if CREDENTIALS_FILE.exists():
        try:
            return json.loads(CREDENTIALS_FILE.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


def _save(creds: dict) -> None:
    CREDENTIALS_FILE.parent.mkdir(parents=True, exist_ok=True)
    CREDENTIALS_FILE.write_text(
        json.dumps(creds, indent=2, ensure_ascii=False), encoding="utf-8"
    )


def _ensure_user(email: str, password: str, password_hash: str | None = None) -> tuple[str, str]:
    """Create the user, or reset their password if they already exist.

    Returns ``(status, password_hash)`` so callers can persist the hash for
    later bootstrap on the deployed Streamlit Cloud database.
    """
    if password_hash is None:
        password_hash = bcrypt.hashpw(
            password.encode("utf-8"), bcrypt.gensalt()
        ).decode("utf-8")
    existing = db.get_user_by_email(email)
    if existing is None:
        ok, msg = db.create_user(email, password_hash, is_admin=False)
        return ("created" if ok else f"create-failed: {msg}", password_hash)
    ok, msg = db.update_user_password(email, password_hash)
    return ("password-reset" if ok else f"reset-failed: {msg}", password_hash)


# ---------------------------------------------------------------------------
# Word document
# ---------------------------------------------------------------------------

def _write_docx(rows: list[dict]) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.table import WD_ALIGN_VERTICAL

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Title
    title = doc.add_heading("ICBB-AMRSS — Sentinel Lab Login Credentials", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x14, 0x3B, 0x5C)

    intro = doc.add_paragraph()
    intro.add_run(
        "Each sentinel-site laboratory listed below has a permanent account on the "
        "ICBB AMR Surveillance System (ICBB-AMRSS). Use the username and password "
        "shown to sign in at "
    )
    link_run = intro.add_run("https://icbb-amrss.streamlit.app")
    link_run.bold = True
    intro.add_run(
        ".\n\nThese credentials are fixed and used for live submissions. Please do "
        "not share them outside your laboratory. Contact the ICBB administrator if "
        "you suspect they have been compromised."
    )

    doc.add_paragraph()

    # Credentials table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Laboratory"
    hdr[1].text = "Username"
    hdr[2].text = "Password"
    for cell in hdr:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["lab"]
        cells[1].text = row["email"]
        cells[2].text = row["password"]
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10.5)

    # Set column widths
    widths = (Cm(7.5), Cm(5.5), Cm(4.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run = note.add_run(
        "Sign-in tips: usernames are case-insensitive; passwords are case-sensitive. "
        "If your password does not work, please contact the ICBB administrator."
    )
    note_run.italic = True
    note_run.font.size = Pt(10)

    doc.save(WORD_OUTPUT)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    creds = _load_existing()
    rows: list[dict] = []
    public: dict[str, dict] = {}
    print(f"Provisioning {len(APPROVED_LABS)} lab accounts...\n")

    for lab_name in APPROVED_LABS:
        code = LAB_CODES.get(lab_name)
        if not code:
            print(f"  ! No short code for '{lab_name}' -- skipped")
            continue
        email = f"{code}@{EMAIL_DOMAIN}"

        entry = creds.get(email) or {}
        password = entry.get("password")
        password_hash = entry.get("password_hash")
        if not password:
            password = _make_password(code)
            password_hash = None  # force fresh hash below
        if not password_hash:
            password_hash = bcrypt.hashpw(
                password.encode("utf-8"), bcrypt.gensalt()
            ).decode("utf-8")

        creds[email] = {
            "lab": lab_name,
            "code": code,
            "password": password,
            "password_hash": password_hash,
        }

        status, _ = _ensure_user(email, password, password_hash)
        print(f"  {lab_name[:48]:48} {email:32} [{status}]")
        rows.append({"lab": lab_name, "email": email, "password": password})
        public[email] = {"lab": lab_name, "code": code, "password_hash": password_hash}

    _save(creds)
    PUBLIC_HASH_FILE.parent.mkdir(parents=True, exist_ok=True)
    PUBLIC_HASH_FILE.write_text(
        json.dumps(public, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    _write_docx(rows)
    print(f"\nWrote credentials JSON     -> {CREDENTIALS_FILE}  (gitignored)")
    print(f"Wrote public hash manifest -> {PUBLIC_HASH_FILE}  (committed)")
    print(f"Wrote Word document        -> {WORD_OUTPUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
